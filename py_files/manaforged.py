# Top level for program
import re
import subprocess
import sys
from docx import Document
from docx.shared import Inches
"""
    Manaforged makes it easier than ever to automatically
    inflate your income and not leave a paper trail..., ok
    Paul?
    It converts PDFs to WordDocs, or vice versa, and increases
    any numbers found in the document.
"""
# converts pdf to text
def pdf_to_text(pdf_name):
    pdf_name = sys.argv[1]
    pdf_content = subprocess.check_output(["pdf2txt.py", pdf_name])
    return pdf_content

# Need to manaforgize the numbers in here
def text_to_word_doc(text):
    document = Document()
    document.add_heading('Document Title', 0)
    lines = text.split('\n')
    # print(lines)

    # cleaning up the text
    for line in text.split('\n'):
        ascii_pdf_content = line.decode("utf-8", "replace")
        control_char_map = dict.fromkeys(range(32))
        clean_pdf_content = ascii_pdf_content.translate(control_char_map)
        #look_for_money(clean_pdf_content)
        document.add_paragraph(clean_pdf_content)
    return document

# Look for money
def look_for_money(pdf_content):
    money = re.findall("\$+\d*[.,]\d*[..]\d*", pdf_content)
    return money

def increase_money(money):
    print(money)
    for item in money:
        item = item.replace("$", "")
        item = item.replace(",", "")
        as_float = float(item)
        as_float *= 2
        print(as_float)


if __name__ == "__main__":
    pdf_text =  pdf_to_text(sys.argv[1])
    money = look_for_money(pdf_text)
    increase_money(money)
    document = text_to_word_doc(pdf_text)
    document.save('demo.docx')
